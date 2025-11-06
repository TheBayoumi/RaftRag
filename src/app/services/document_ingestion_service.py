"""
Custom Document Ingestion Service with BM25 + Semantic Hybrid Retrieval.

This service implements hybrid search combining:
- BM25 (keyword-based, statistical)
- Semantic search (embedding-based, neural)
- RRF (Reciprocal Rank Fusion) for score combination

This approach provides:
- Better precision (exact keyword matching)
- Better recall (semantic similarity)
- Industry-standard retrieval (Elasticsearch/OpenSearch use this)
"""

import pickle
import uuid
from collections import defaultdict
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import chromadb
import numpy as np
import pandas as pd
from chromadb.config import Settings as ChromaSettings
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from loguru import logger
from pypdf import PdfReader
from rank_bm25 import BM25Okapi

from ..core.config import get_settings
from ..core.exceptions import RAGException
from ..schemas.rag import DocumentResponse
from ..utils.local_models import LocalEmbeddingWrapper
from .base import BaseService

settings = get_settings()


class DocumentIngestionService(BaseService):
    """
    Hybrid document ingestion service with BM25 + semantic search.

    Features:
    - ChromaDB for semantic (vector) search
    - BM25 for keyword-based search
    - RRF for combining both scores
    - Automatic index persistence
    """

    def __init__(self) -> None:
        """Initialize hybrid document ingestion service."""
        super().__init__("DocumentIngestionService")

        # Semantic search components
        self.embedding_wrapper: Optional[LocalEmbeddingWrapper] = None
        self.chroma_clients: Dict[str, chromadb.Client] = {}
        self.collections: Dict[str, chromadb.Collection] = {}
        self.text_splitter: Optional[RecursiveCharacterTextSplitter] = None

        # BM25 components (per collection)
        self.bm25_indices: Dict[str, BM25Okapi] = {}  # collection_name -> BM25 index
        self.bm25_documents: Dict[str, List[str]] = {}  # collection_name -> doc texts
        self.bm25_metadata: Dict[str, List[Dict]] = (
            {}
        )  # collection_name -> doc metadata

        # BM25 index storage paths (from settings)
        self.bm25_base_dir = settings.bm25_indices_dir
        # Directory is created in ensure_directories() on startup

    async def _initialize_impl(self) -> None:
        """
        Initialize hybrid document ingestion service.

        Returns:
            None
        """
        # Initialize embedding model
        self.logger.info("Initializing embedding model")
        self.embedding_wrapper = LocalEmbeddingWrapper()
        self.embedding_wrapper.load_model()
        self.logger.success("Embedding model initialized")

        # Initialize text splitter
        self.logger.info("Initializing RecursiveCharacterTextSplitter")
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.default_chunk_size,
            chunk_overlap=settings.default_chunk_overlap,
            length_function=len,
        )
        self.logger.success("Text splitter initialized")

        # Initialize ChromaDB clients for RAFT and RAG
        await self._initialize_chromadb_clients()

        # Load existing BM25 indices
        await self._load_bm25_indices()

        # Build BM25 indices for collections that don't have them
        await self._build_missing_bm25_indices()

    async def _initialize_chromadb_clients(self) -> None:
        """Initialize separate ChromaDB clients for RAFT and RAG databases."""
        # RAFT database (for fine-tuning dataset generation)
        raft_persist_dir = Path(settings.chroma_persist_dir) / "raft"
        raft_persist_dir.mkdir(parents=True, exist_ok=True)

        self.chroma_clients["raft"] = chromadb.Client(
            ChromaSettings(
                anonymized_telemetry=False,
                allow_reset=True,
                is_persistent=True,
                persist_directory=str(raft_persist_dir),
            )
        )
        self.logger.info(f"RAFT ChromaDB initialized at {raft_persist_dir}")

        # RAG database (for production retrieval)
        rag_persist_dir = Path(settings.chroma_persist_dir) / "rag"
        rag_persist_dir.mkdir(parents=True, exist_ok=True)

        self.chroma_clients["rag"] = chromadb.Client(
            ChromaSettings(
                anonymized_telemetry=False,
                allow_reset=True,
                is_persistent=True,
                persist_directory=str(rag_persist_dir),
            )
        )
        self.logger.info(f"RAG ChromaDB initialized at {rag_persist_dir}")

    async def _load_bm25_indices(self) -> None:
        """Load existing BM25 indices from disk."""
        for db_type in ["raft", "rag"]:
            db_dir = self.bm25_base_dir / db_type
            if not db_dir.exists():
                continue

            for index_file in db_dir.glob("*.pkl"):
                collection_name = index_file.stem
                try:
                    with open(index_file, "rb") as f:
                        data = pickle.load(f)

                    key = f"{db_type}_{collection_name}"
                    self.bm25_indices[key] = data["index"]
                    self.bm25_documents[key] = data["documents"]
                    self.bm25_metadata[key] = data["metadata"]

                    self.logger.info(
                        f"Loaded BM25 index for {db_type}/{collection_name} "
                        f"({len(data['documents'])} documents)"
                    )
                except Exception as e:
                    self.logger.warning(f"Failed to load BM25 index {index_file}: {e}")

    async def _build_missing_bm25_indices(self) -> None:
        """
        Build BM25 indices for existing ChromaDB collections that don't have them.

        This is called during initialization to ensure all collections have BM25 indices.
        """
        for db_type in ["raft", "rag"]:
            client = self.chroma_clients.get(db_type)
            if not client:
                continue

            try:
                collections = client.list_collections()
                self.logger.info(
                    f"Checking BM25 indices for {len(collections)} collections "
                    f"in {db_type} database"
                )

                for collection in collections:
                    collection_name = collection.name
                    key = f"{db_type}_{collection_name}"

                    # Skip if BM25 index already exists
                    if key in self.bm25_indices:
                        self.logger.debug(
                            f"BM25 index already exists for {db_type}/{collection_name}"
                        )
                        continue

                    self.logger.info(
                        f"Building BM25 index for {db_type}/{collection_name}..."
                    )

                    # Get all documents from ChromaDB
                    results = collection.get(include=["documents", "metadatas"])

                    if not results["documents"]:
                        self.logger.debug(f"No documents in {collection_name}")
                        continue

                    documents = results["documents"]
                    metadatas = results["metadatas"]

                    self.logger.info(
                        f"Found {len(documents)} document chunks in {collection_name}"
                    )

                    # Initialize BM25 storage
                    self.bm25_documents[key] = documents
                    self.bm25_metadata[key] = metadatas

                    # Build BM25 index
                    tokenized_corpus = [self._tokenize(doc) for doc in documents]
                    self.bm25_indices[key] = BM25Okapi(tokenized_corpus)

                    self.logger.success(
                        f"Built BM25 index for {db_type}/{collection_name} "
                        f"with {len(documents)} documents"
                    )

                    # Save to disk
                    self._save_bm25_index(collection_name, db_type)

            except Exception as e:
                self.logger.error(
                    f"Failed to build BM25 indices for {db_type} database: {e}"
                )

    def _save_bm25_index(self, collection_name: str, database_type: str) -> None:
        """
        Save BM25 index to disk.

        Args:
            collection_name: Collection name.
            database_type: Either "raft" or "rag".
        """
        key = f"{database_type}_{collection_name}"
        if key not in self.bm25_indices:
            return

        db_dir = self.bm25_base_dir / database_type
        db_dir.mkdir(parents=True, exist_ok=True)

        index_file = db_dir / f"{collection_name}.pkl"

        try:
            data = {
                "index": self.bm25_indices[key],
                "documents": self.bm25_documents[key],
                "metadata": self.bm25_metadata[key],
            }

            with open(index_file, "wb") as f:
                pickle.dump(data, f)

            self.logger.debug(f"Saved BM25 index to {index_file}")
        except Exception as e:
            self.logger.error(f"Failed to save BM25 index: {e}")

    def _tokenize(self, text: str) -> List[str]:
        """
        Tokenize text for BM25.

        Simple whitespace + lowercase tokenization.
        For better results, could use nltk or spacy.

        Args:
            text: Input text.

        Returns:
            List[str]: Tokens.
        """
        return text.lower().split()

    def _get_or_create_collection(
        self, collection_name: str, database_type: str = "rag"
    ) -> chromadb.Collection:
        """
        Get or create a ChromaDB collection.

        Args:
            collection_name: Collection name.
            database_type: Either "raft" or "rag".

        Returns:
            chromadb.Collection: Collection instance.
        """
        collection_key = f"{database_type}_{collection_name}"

        if collection_key in self.collections:
            return self.collections[collection_key]

        client = self.chroma_clients.get(database_type)
        if not client:
            raise RAGException(
                f"No ChromaDB client for database type: {database_type}",
                operation="get_collection",
            )

        collection = client.get_or_create_collection(
            name=collection_name,
            metadata={"hnsw:space": "cosine"},
        )

        self.collections[collection_key] = collection
        self.logger.info(
            f"Collection '{collection_name}' ready in {database_type} database"
        )

        return collection

    def _parse_text_file(self, file_path: Path) -> str:
        """Parse plain text file."""
        try:
            return file_path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            return file_path.read_text(encoding="latin-1")

    def _parse_pdf_file(self, file_path: Path) -> str:
        """Parse PDF file."""
        reader = PdfReader(str(file_path))
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text

    def _parse_csv_file(self, file_path: Path) -> str:
        """
        Parse CSV file and convert to text format.

        Converts CSV rows to a readable text format where each row is represented
        as "Column1: value1, Column2: value2, ..." for better semantic search.

        Args:
            file_path: Path to CSV file.

        Returns:
            str: Text representation of CSV data.

        Raises:
            RAGException: If CSV parsing fails.
        """
        try:
            # Try reading with different encodings
            try:
                df = pd.read_csv(file_path, encoding="utf-8")
            except UnicodeDecodeError:
                df = pd.read_csv(file_path, encoding="latin-1")

            # Handle empty CSV
            if df.empty:
                return ""

            # Convert each row to a text representation
            # Format: "Column1: value1, Column2: value2, ..."
            rows = []
            for idx, row in df.iterrows():
                # Convert row to dictionary and format as text
                row_text = ", ".join(
                    f"{col}: {val}" for col, val in row.items() if pd.notna(val)
                )
                if row_text:  # Only add non-empty rows
                    rows.append(f"Row {idx + 1}: {row_text}")

            # Join all rows with newlines
            text = "\n".join(rows)

            # If no rows were added, include column headers at least
            if not text:
                headers = ", ".join(df.columns.tolist())
                text = f"Columns: {headers}"

            return text

        except Exception as e:
            raise RAGException(
                f"Failed to parse CSV file {file_path.name}: {e}",
                operation="parse_csv",
            )

    def _parse_docx_file(self, file_path: Path) -> str:
        """
        Parse Word document (.docx) file and extract text.

        Extracts text from all paragraphs and tables in the document,
        preserving paragraph structure with newlines.

        Note: Only .docx format is supported (not legacy .doc format).

        Args:
            file_path: Path to Word document file (.docx).

        Returns:
            str: Extracted text from the document.

        Raises:
            RAGException: If Word document parsing fails.
        """
        try:
            doc = DocxDocument(str(file_path))
            paragraphs = []

            # Extract text from all paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:  # Only add non-empty paragraphs
                    paragraphs.append(text)

            # Extract text from tables if present
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        paragraphs.append(row_text)

            # Join paragraphs with newlines
            text = "\n".join(paragraphs)

            return text

        except Exception as e:
            raise RAGException(
                f"Failed to parse Word document {file_path.name}: {e}",
                operation="parse_docx",
            )

    async def ingest_document(
        self,
        file_path: str,
        collection_name: str = "default",
        chunk_size: int = 512,
        chunk_overlap: int = 50,
        metadata: Optional[Dict[str, Any]] = None,
        database_type: str = "rag",
    ) -> DocumentResponse:
        """
        Ingest document with hybrid indexing (BM25 + semantic).

        Args:
            file_path: Path to document file.
            collection_name: Collection name.
            chunk_size: Chunk size for splitting.
            chunk_overlap: Overlap between chunks.
            metadata: Optional metadata.
            database_type: Either "raft" or "rag".

        Returns:
            DocumentResponse: Ingestion result.

        Raises:
            RAGException: If ingestion fails.
        """
        try:
            path = Path(file_path)
            if not path.exists():
                raise RAGException(f"File not found: {file_path}", operation="ingest")

            self.logger.info(f"Ingesting document: {path.name}")

            # Parse document based on file type
            if path.suffix.lower() in [".txt", ".md"]:
                text = self._parse_text_file(path)
            elif path.suffix.lower() == ".pdf":
                text = self._parse_pdf_file(path)
            elif path.suffix.lower() == ".csv":
                text = self._parse_csv_file(path)
            elif path.suffix.lower() == ".docx":
                text = self._parse_docx_file(path)
            else:
                raise RAGException(
                    f"Unsupported file type: {path.suffix}",
                    operation="ingest",
                )

            if not text or len(text.strip()) == 0:
                raise RAGException(
                    f"No text extracted from {path.name}", operation="ingest"
                )

            # Split into chunks
            self.text_splitter.chunk_size = chunk_size
            self.text_splitter.chunk_overlap = chunk_overlap
            chunks = self.text_splitter.split_text(text)

            if not chunks:
                raise RAGException(
                    f"No chunks created from {path.name}", operation="ingest"
                )

            self.logger.info(f"Created {len(chunks)} chunks from {path.name}")

            # Generate document ID
            doc_id = str(uuid.uuid4())

            # Prepare metadata for all chunks
            base_metadata = {
                "filename": path.name,
                "file_type": path.suffix,
                "doc_id": doc_id,
                "collection": collection_name,
                "database_type": database_type,
                "total_chunks": len(chunks),
                "created_at": datetime.now().isoformat(),
            }
            if metadata:
                base_metadata.update(metadata)

            # Get/create collection
            collection = self._get_or_create_collection(collection_name, database_type)

            # Generate embeddings for semantic search
            chunk_embeddings = self.embedding_wrapper.embed_documents(chunks)

            # Prepare data for ChromaDB
            chunk_ids = [f"{doc_id}_chunk_{i}" for i in range(len(chunks))]
            chunk_metadatas = [
                {**base_metadata, "chunk_index": i} for i in range(len(chunks))
            ]

            # Store in ChromaDB (semantic index)
            collection.add(
                ids=chunk_ids,
                embeddings=chunk_embeddings,
                documents=chunks,
                metadatas=chunk_metadatas,
            )

            self.logger.info(f"Stored {len(chunks)} chunks in ChromaDB")

            # Build/update BM25 index (keyword index)
            key = f"{database_type}_{collection_name}"

            # Initialize BM25 index if not exists
            if key not in self.bm25_documents:
                self.bm25_documents[key] = []
                self.bm25_metadata[key] = []

            # Add documents to BM25 corpus
            self.bm25_documents[key].extend(chunks)
            self.bm25_metadata[key].extend(chunk_metadatas)

            # Rebuild BM25 index with all documents
            tokenized_corpus = [self._tokenize(doc) for doc in self.bm25_documents[key]]
            self.bm25_indices[key] = BM25Okapi(tokenized_corpus)

            self.logger.info(
                f"Updated BM25 index: {len(self.bm25_documents[key])} total documents"
            )

            # Save BM25 index to disk
            self._save_bm25_index(collection_name, database_type)

            # Calculate file size
            file_size_mb = path.stat().st_size / (1024 * 1024)

            return DocumentResponse(
                doc_id=doc_id,
                filename=path.name,
                collection_name=collection_name,
                num_chunks=len(chunks),
                file_size_mb=round(file_size_mb, 2),
                created_at=datetime.now(),
                status="completed",
            )

        except Exception as e:
            self.handle_error(e, {"file": file_path})
            raise RAGException(f"Failed to ingest document: {e}", operation="ingest")

    def _reciprocal_rank_fusion(
        self,
        semantic_results: List[Tuple[str, float, Dict]],
        bm25_results: List[Tuple[str, float, Dict]],
        k: int = 60,
    ) -> List[Dict[str, Any]]:
        """
        Combine semantic and BM25 results using Reciprocal Rank Fusion.

        RRF formula: score = 1/(k + rank)
        Scores are normalized to 0-1 range for compatibility with similarity thresholds.

        Args:
            semantic_results: List of (doc_text, score, metadata) from semantic search.
            bm25_results: List of (doc_text, score, metadata) from BM25.
            k: RRF constant (default 60, standard value).

        Returns:
            List[Dict]: Combined results sorted by normalized RRF score (0-1 range).
        """
        # Calculate maximum possible RRF score (when document ranks #1 in both)
        # max_rrf = 1/(k+1) + 1/(k+1) = 2/(k+1)
        max_rrf = 2.0 / (k + 1)

        # Build score dictionaries (use doc_text as key for matching)
        semantic_scores = {}
        bm25_scores = {}
        all_docs = {}  # doc_text -> (original_score, metadata)

        # Process semantic results (rank 1, 2, 3...)
        for rank, (doc_text, score, metadata) in enumerate(semantic_results, 1):
            rrf_score = 1.0 / (k + rank)
            semantic_scores[doc_text] = rrf_score
            all_docs[doc_text] = (score, metadata)

        # Process BM25 results (rank 1, 2, 3...)
        for rank, (doc_text, score, metadata) in enumerate(bm25_results, 1):
            rrf_score = 1.0 / (k + rank)
            bm25_scores[doc_text] = rrf_score
            if doc_text not in all_docs:
                all_docs[doc_text] = (score, metadata)

        # Combine scores and normalize to 0-1 range
        combined = []
        for doc_text, (original_score, metadata) in all_docs.items():
            # Raw RRF score
            raw_rrf = semantic_scores.get(doc_text, 0) + bm25_scores.get(doc_text, 0)

            # Normalize to 0-1 range (divide by max possible score)
            normalized_score = raw_rrf / max_rrf

            combined.append(
                {
                    "content": doc_text,
                    "score": normalized_score,  # Normalized RRF score (0-1 range)
                    "raw_rrf_score": raw_rrf,  # Keep raw RRF for debugging
                    "semantic_score": semantic_scores.get(doc_text, 0),
                    "bm25_score": bm25_scores.get(doc_text, 0),
                    "metadata": metadata,
                    "doc_id": metadata.get("doc_id", "unknown"),
                }
            )

        # Sort by normalized RRF score (highest first)
        combined.sort(key=lambda x: x["score"], reverse=True)

        return combined

    async def query_documents(
        self,
        query: str,
        collection_name: str,
        database_type: str = "rag",
        top_k: int = 5,
        similarity_threshold: float = 0.3,
    ) -> List[Dict[str, Any]]:
        """
        Hybrid query using BM25 + semantic search with RRF fusion.

        Args:
            query: Query text.
            collection_name: Collection to search.
            database_type: Either "raft" or "rag".
            top_k: Number of results to return.
            similarity_threshold: Minimum RRF score threshold.

        Returns:
            List[Dict]: Retrieved documents with RRF scores.

        Raises:
            RAGException: If query fails.
        """
        try:
            collection = self._get_or_create_collection(collection_name, database_type)
            key = f"{database_type}_{collection_name}"

            # 1. Semantic search
            query_embedding = self.embedding_wrapper.embed_query(query)

            if hasattr(query_embedding, "tolist"):
                query_embedding_list = query_embedding.tolist()
            else:
                query_embedding_list = query_embedding

            # Retrieve more candidates for better fusion (2x top_k)
            semantic_k = top_k * 2

            semantic_raw = collection.query(
                query_embeddings=[query_embedding_list], n_results=semantic_k
            )

            semantic_results = []
            for i, (doc, metadata, distance) in enumerate(
                zip(
                    semantic_raw["documents"][0],
                    semantic_raw["metadatas"][0],
                    semantic_raw["distances"][0],
                )
            ):
                # Convert cosine distance to similarity
                similarity = 1.0 - distance
                semantic_results.append((doc, similarity, metadata))

            # 2. BM25 search
            bm25_results = []
            if key in self.bm25_indices:
                tokenized_query = self._tokenize(query)
                bm25_scores = self.bm25_indices[key].get_scores(tokenized_query)

                # Get top semantic_k results
                top_indices = np.argsort(bm25_scores)[::-1][:semantic_k]

                for idx in top_indices:
                    if idx < len(self.bm25_documents[key]):
                        doc = self.bm25_documents[key][idx]
                        score = float(bm25_scores[idx])
                        metadata = self.bm25_metadata[key][idx]
                        bm25_results.append((doc, score, metadata))
            else:
                self.logger.warning(
                    f"No BM25 index for {database_type}/{collection_name}, "
                    "using semantic-only search"
                )

            # 3. Combine with RRF
            if bm25_results:
                combined_results = self._reciprocal_rank_fusion(
                    semantic_results, bm25_results, k=60
                )

                # Log hybrid search statistics
                if combined_results:
                    top_score = combined_results[0]["score"]
                    min_score = combined_results[-1]["score"]
                    self.logger.debug(
                        f"Hybrid search: {len(semantic_results)} semantic + "
                        f"{len(bm25_results)} BM25 → {len(combined_results)} combined "
                        f"(scores: {top_score:.3f} - {min_score:.3f})"
                    )
                else:
                    self.logger.debug(
                        f"Hybrid search: {len(semantic_results)} semantic + "
                        f"{len(bm25_results)} BM25 → 0 combined"
                    )
            else:
                # Fallback to semantic only if no BM25 index
                combined_results = [
                    {
                        "content": doc,
                        "score": score,
                        "semantic_score": score,
                        "bm25_score": 0.0,
                        "metadata": metadata,
                        "doc_id": metadata.get("doc_id", "unknown"),
                    }
                    for doc, score, metadata in semantic_results
                ]

            # 4. Filter by threshold
            original_count = len(combined_results)
            filtered_docs = [
                doc
                for doc in combined_results[:top_k]
                if doc["score"] >= similarity_threshold
            ]

            self.logger.debug(
                f"Retrieved {original_count} documents, "
                f"{len(filtered_docs)} passed threshold {similarity_threshold:.2f} "
                f"(top-{top_k} filtered to {len(filtered_docs)})"
            )

            return filtered_docs

        except Exception as e:
            self.handle_error(e, {"query": query, "collection": collection_name})
            raise RAGException(f"Failed to query documents: {e}", operation="query")

    async def delete_document(
        self, doc_id: str, collection_name: str, database_type: str = "rag"
    ) -> None:
        """
        Delete document from both ChromaDB and BM25 index.

        Args:
            doc_id: Document identifier.
            collection_name: Collection name.
            database_type: Either "raft" or "rag".

        Raises:
            RAGException: If deletion fails.
        """
        try:
            collection = self._get_or_create_collection(collection_name, database_type)
            key = f"{database_type}_{collection_name}"

            # Get all chunk IDs for this document
            results = collection.get(where={"doc_id": doc_id})
            chunk_ids = results["ids"]

            if not chunk_ids:
                self.logger.warning(f"No chunks found for document {doc_id}")
                return

            # Delete from ChromaDB
            collection.delete(ids=chunk_ids)
            self.logger.info(f"Deleted {len(chunk_ids)} chunks from ChromaDB")

            # Delete from BM25 index
            if key in self.bm25_documents:
                indices_to_remove = []
                for i, metadata in enumerate(self.bm25_metadata[key]):
                    if metadata.get("doc_id") == doc_id:
                        indices_to_remove.append(i)

                # Remove in reverse order to maintain indices
                for idx in reversed(indices_to_remove):
                    del self.bm25_documents[key][idx]
                    del self.bm25_metadata[key][idx]

                # Rebuild BM25 index
                if self.bm25_documents[key]:
                    tokenized_corpus = [
                        self._tokenize(doc) for doc in self.bm25_documents[key]
                    ]
                    self.bm25_indices[key] = BM25Okapi(tokenized_corpus)
                else:
                    # Remove empty index
                    del self.bm25_indices[key]
                    del self.bm25_documents[key]
                    del self.bm25_metadata[key]

                # Save updated index
                self._save_bm25_index(collection_name, database_type)

                self.logger.info(
                    f"Deleted {len(indices_to_remove)} chunks from BM25 index"
                )

        except Exception as e:
            self.handle_error(e, {"doc_id": doc_id, "collection": collection_name})
            raise RAGException(
                f"Failed to delete document: {e}", operation="delete_document"
            )

    async def ingest_document_from_text(
        self,
        text: str,
        collection_name: str = "default",
        database_type: str = "rag",
        metadata: Optional[Dict[str, Any]] = None,
        chunk_size: int = 512,
        chunk_overlap: int = 50,
    ) -> DocumentResponse:
        """
        Ingest text directly without file (for raganything integration).

        This method is used by ChromaVectorDBStorage adapter to ingest
        documents from raganything.

        Args:
            text: Text content to ingest.
            collection_name: Collection name.
            database_type: Either "raft" or "rag".
            metadata: Optional metadata dictionary.
            chunk_size: Chunk size for splitting.
            chunk_overlap: Overlap between chunks.

        Returns:
            DocumentResponse: Ingestion result.

        Raises:
            RAGException: If ingestion fails.
        """
        try:
            if not text or len(text.strip()) == 0:
                raise RAGException("Empty text provided", operation="ingest")

            self.logger.info(f"Ingesting text ({len(text)} chars)")

            # Split into chunks
            self.text_splitter.chunk_size = chunk_size
            self.text_splitter.chunk_overlap = chunk_overlap
            chunks = self.text_splitter.split_text(text)

            if not chunks:
                raise RAGException("No chunks created from text", operation="ingest")

            self.logger.info(f"Created {len(chunks)} chunks")

            # Generate document ID
            doc_id = str(uuid.uuid4())

            # Prepare metadata for all chunks
            base_metadata = {
                "doc_id": doc_id,
                "collection": collection_name,
                "database_type": database_type,
                "total_chunks": len(chunks),
                "created_at": datetime.now().isoformat(),
                "source": "text_ingestion",
            }
            if metadata:
                base_metadata.update(metadata)

            # Get/create collection
            collection = self._get_or_create_collection(collection_name, database_type)

            # Generate embeddings for semantic search
            chunk_embeddings = self.embedding_wrapper.embed_documents(chunks)

            # Prepare data for ChromaDB
            chunk_ids = [f"{doc_id}_chunk_{i}" for i in range(len(chunks))]
            chunk_metadatas = [
                {**base_metadata, "chunk_index": i} for i in range(len(chunks))
            ]

            # Store in ChromaDB (semantic index)
            collection.add(
                ids=chunk_ids,
                embeddings=chunk_embeddings,
                documents=chunks,
                metadatas=chunk_metadatas,
            )

            self.logger.info(f"Stored {len(chunks)} chunks in ChromaDB")

            # Build/update BM25 index (keyword index)
            key = f"{database_type}_{collection_name}"

            # Initialize BM25 index if not exists
            if key not in self.bm25_documents:
                self.bm25_documents[key] = []
                self.bm25_metadata[key] = []

            # Add documents to BM25 corpus
            self.bm25_documents[key].extend(chunks)
            self.bm25_metadata[key].extend(chunk_metadatas)

            # Rebuild BM25 index with all documents
            tokenized_corpus = [self._tokenize(doc) for doc in self.bm25_documents[key]]
            self.bm25_indices[key] = BM25Okapi(tokenized_corpus)

            self.logger.info(
                f"Updated BM25 index: {len(self.bm25_documents[key])} total documents"
            )

            # Save BM25 index to disk
            self._save_bm25_index(collection_name, database_type)

            return DocumentResponse(
                doc_id=doc_id,
                filename="text_document",
                collection_name=collection_name,
                num_chunks=len(chunks),
                file_size_mb=round(len(text) / (1024 * 1024), 2),
                created_at=datetime.now(),
                status="completed",
            )

        except Exception as error:
            self.handle_error(error, {"text_length": len(text)})
            raise RAGException(f"Failed to ingest text: {error}", operation="ingest")

    async def delete_document_by_metadata(
        self,
        collection_name: str,
        database_type: str,
        metadata_filter: Dict[str, Any],
    ) -> int:
        """
        Delete documents by metadata filter.

        Args:
            collection_name: Collection name.
            database_type: Either "raft" or "rag".
            metadata_filter: Metadata filter dictionary.

        Returns:
            int: Number of documents deleted.

        Raises:
            RAGException: If deletion fails.
        """
        try:
            collection = self._get_or_create_collection(collection_name, database_type)
            key = f"{database_type}_{collection_name}"

            # Query documents matching filter
            results = collection.get(where=metadata_filter)
            chunk_ids = results["ids"]

            if not chunk_ids:
                self.logger.warning(
                    f"No documents found matching filter: {metadata_filter}"
                )
                return 0

            # Delete from ChromaDB
            collection.delete(ids=chunk_ids)
            self.logger.info(f"Deleted {len(chunk_ids)} chunks from ChromaDB")

            # Delete from BM25 index
            if key in self.bm25_documents:
                indices_to_remove = []
                for i, doc_metadata in enumerate(self.bm25_metadata[key]):
                    # Check if all filter conditions match
                    if all(
                        doc_metadata.get(k) == v for k, v in metadata_filter.items()
                    ):
                        indices_to_remove.append(i)

                # Remove in reverse order to maintain indices
                for idx in reversed(indices_to_remove):
                    del self.bm25_documents[key][idx]
                    del self.bm25_metadata[key][idx]

                # Rebuild BM25 index if documents remain
                if self.bm25_documents[key]:
                    tokenized_corpus = [
                        self._tokenize(doc) for doc in self.bm25_documents[key]
                    ]
                    self.bm25_indices[key] = BM25Okapi(tokenized_corpus)
                else:
                    # Remove empty index
                    del self.bm25_indices[key]
                    del self.bm25_documents[key]
                    del self.bm25_metadata[key]

                # Save updated index
                self._save_bm25_index(collection_name, database_type)

                self.logger.info(
                    f"Deleted {len(indices_to_remove)} chunks from BM25 index"
                )

            return len(chunk_ids)

        except Exception as error:
            self.handle_error(
                error, {"collection": collection_name, "filter": metadata_filter}
            )
            raise RAGException(
                f"Failed to delete by metadata: {error}",
                operation="delete_by_metadata",
            )

    async def delete_collection(
        self, collection_name: str, database_type: str = "rag"
    ) -> None:
        """
        Delete entire collection from both ChromaDB and BM25.

        Args:
            collection_name: Collection name to delete.
            database_type: Either "raft" or "rag".

        Raises:
            RAGException: If deletion fails.
        """
        try:
            key = f"{database_type}_{collection_name}"

            # Delete from ChromaDB
            client = self.chroma_clients.get(database_type)
            if client:
                try:
                    client.delete_collection(name=collection_name)
                    self.logger.info(f"Deleted ChromaDB collection: {collection_name}")
                except Exception as chroma_error:
                    self.logger.warning(
                        f"ChromaDB collection not found or already deleted: "
                        f"{chroma_error}"
                    )

            # Remove from local collections dict
            if key in self.collections:
                del self.collections[key]

            # Delete BM25 index
            if key in self.bm25_indices:
                del self.bm25_indices[key]
            if key in self.bm25_documents:
                del self.bm25_documents[key]
            if key in self.bm25_metadata:
                del self.bm25_metadata[key]

            # Delete persisted BM25 index file
            bm25_file = self.bm25_base_dir / database_type / f"{collection_name}.pkl"
            if bm25_file.exists():
                bm25_file.unlink()
                self.logger.info(f"Deleted BM25 index file: {bm25_file}")

            self.logger.success(f"Collection deleted: {collection_name}")

        except Exception as error:
            self.handle_error(
                error, {"collection": collection_name, "database": database_type}
            )
            raise RAGException(
                f"Failed to delete collection: {error}",
                operation="delete_collection",
            )

    def __getstate__(self) -> Dict[str, Any]:
        """
        Prepare service for pickling.

        Removes non-picklable objects (ChromaDB clients, text splitter)
        before serialization.

        Returns:
            Dict[str, Any]: Picklable state dictionary.
        """
        state = super().__getstate__()

        # Remove ChromaDB clients (contain non-picklable state)
        state.pop("chroma_clients", None)
        state.pop("collections", None)

        # Remove text splitter (can be recreated)
        state.pop("text_splitter", None)

        # Remove embedding wrapper reference (will be recreated on initialization)
        state.pop("embedding_wrapper", None)

        # BM25 indices, documents, and metadata can be pickled
        # They will persist across pickle/unpickle

        return state

    def __setstate__(self, state: Dict[str, Any]) -> None:
        """
        Restore service after unpickling.

        Reconstructs non-picklable objects.

        Args:
            state: State dictionary from pickle.
        """
        super().__setstate__(state)

        # Set non-picklable objects to None (will be recreated on demand)
        self.chroma_clients = {}
        self.collections = {}
        self.text_splitter = None
        self.embedding_wrapper = None

        # Mark as uninitialized to ensure reinitialization
        self._initialized = False
        self.logger.warning(
            f"{self.service_name} was unpickled and needs reinitialization. "
            "Call initialize() before using."
        )
